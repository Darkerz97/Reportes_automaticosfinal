import streamlit as st 
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


#create a new document (instacia el documento para crear un documento en blanco)
doc = Document()

#titulo de la app
st.write("Generador de reportes")


# Initialize connection. establece conexion usando las credenciales de secrets
conn = st.connection('mysql', type='sql')


#query a la base de datos
df = conn.query('select * from reparaciones_limpia;', ttl=600)

#print results imprime resultados de la busqueda en la base 
sbfolios= df['efolio']


#Crea un menu desplegable con los folio en la base de dato
Folio_selection= st.selectbox(
    'seleccione el folio para visualizar la informacion',
    (sbfolios))

#imprime los datos de el folio seleccionado
st.write(df[df['efolio']==Folio_selection])   




#crea el documento y le agrega la informacion 
# Add a title
title = doc.add_heading(f'Reporte del Equipo {Folio_selection} ', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a paragraph with bold and italic text
paragraph = doc.add_paragraph('reporte de diganostico y reparacion')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = paragraph.runs[0]
run.bold = True
run.italic = True

# Add a heading
doc.add_heading('diagnostico', level=2)

# Add a bulleted list
list_paragraph = doc.add_paragraph()
list_paragraph.add_run('Bullet 1').bold = True
list_paragraph.add_run(' - This is the first bullet point.')
list_paragraph.add_run('\n')
list_paragraph.add_run('Bullet 2').bold = True
list_paragraph.add_run(' - This is the second bullet point.')


#guarda el documento
doc.save('demo.docx')




#boton descargable para el documento
with open("demo.docx","rb") as file:
    st.download_button(
        label="Descargar Reporte",
        data=file,
        file_name= (f"{Folio_selection} reporte.docx"),
        mime="demo/docx",
        icon=":material/download:",
)